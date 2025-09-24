# article/views.py
from datetime import datetime, timedelta
import json
import re

from django.contrib.auth import get_user_model
from django.contrib.contenttypes.models import ContentType
from django.db import IntegrityError
from django.db.models import Count, Q
from django.http import JsonResponse, HttpResponseBadRequest
from django.shortcuts import get_object_or_404, redirect
from django.utils.text import slugify
from django.utils.timezone import now
from django.views.decorators.csrf import csrf_exempt
from django.views.generic import ListView, DetailView

from taggit.models import Tag, TaggedItem

from .models import Article, ArticleTranslation

from io import BytesIO

from django.http import HttpResponse
from django.shortcuts import render
from django.utils import timezone
from django.utils.html import strip_tags
from django.db import transaction

from urllib.parse import urlparse, parse_qs


# -----------------------------
# Language helpers
# -----------------------------
SUPPORTED_LANGS = {"ja", "zh", "en"}

def _get_lang(request):
    """Get session language; default ja."""
    lang = request.session.get("lang", "ja")
    return lang if lang in SUPPORTED_LANGS else "ja"


# -----------------------------
# Article views
# -----------------------------
class ArticleDetailView(DetailView):
    model = Article
    template_name = 'article/article_detail.html'
    context_object_name = 'article'

    def get_object(self, queryset=None):
        # Support URLs with year/month/day/slug
        year = self.kwargs.get("year")
        month = self.kwargs.get("month")
        day = self.kwargs.get("day")
        slug = self.kwargs.get("slug")
        if year and month and day and slug:
            return get_object_or_404(
                Article,
                publish__year=year,
                publish__month=month,
                publish__day=day,
                slug=slug
            )
        return super().get_object(queryset)

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        request = self.request
        lang = _get_lang(request)
        article: Article = ctx["article"]

        display_title = article.title
        display_body = article.text

        if lang != "ja":
            tr = ArticleTranslation.objects.filter(article=article, language=lang).first()
            if not tr and request.user.is_authenticated and request.user.is_staff:
                # Staff-only auto-translation
                try:
                    t_title, t_body = translate_title_body(article.title, article.text, lang)
                    tr = ArticleTranslation.objects.create(
                        article=article,
                        language=lang,
                        title_translated=t_title,
                        text_translated=t_body
                    )
                except Exception:
                    tr = None
            if tr:
                display_title = tr.title_translated or display_title
                display_body = tr.text_translated or display_body

        ctx.update({
            "lang": lang,
            "display_title": display_title,
            "display_body": display_body,
        })
        return ctx


class ArticleListView(ListView):
    model = Article
    paginate_by = 15
    template_name = 'article/article_list.html'
    context_object_name = 'articles'

    def get_queryset(self):
        queryset = Article.objects.all().select_related('user').prefetch_related('tags')

        # 🔍 search
        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(
                Q(title__icontains=query) | Q(text__icontains=query)
            )

        # tag filter
        tag_slug = self.request.GET.get('tag')
        if tag_slug:
            queryset = queryset.filter(tags__slug=tag_slug)

        # date filters
        period = self.request.GET.get('period')
        today = now().date()
        if period == 'today':
            queryset = queryset.filter(publish__date=today)
        elif period == '7days':
            queryset = queryset.filter(publish__gte=today - timedelta(days=7))
        elif period == 'month':
            queryset = queryset.filter(publish__year=today.year, publish__month=today.month)
        elif period == 'year':
            queryset = queryset.filter(publish__year=today.year)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        request = self.request
        lang = _get_lang(request)
        today = now().date()
        content_type = ContentType.objects.get_for_model(Article)

        # only tags linked to Article
        tag_ids = TaggedItem.objects.filter(
            content_type=content_type
        ).values_list('tag_id', flat=True)

        tag_list = Tag.objects.filter(id__in=tag_ids).annotate(
            article_count=Count('taggit_taggeditem_items')
        )
        context['tag_list'] = tag_list

        # sidebar date filters
        context['publish_filters'] = [
            ('',     '- Any time', Article.objects.count()),
            ('7days','- Past 7 days', Article.objects.filter(publish__gte=today - timedelta(days=7)).count()),
            ('month','- This month',     Article.objects.filter(publish__year=today.year, publish__month=today.month).count()),
            ('year', '- This year',     Article.objects.filter(publish__year=today.year).count()),
        ]

        context['selected_tag'] = request.GET.get('tag')
        context['selected_period'] = request.GET.get('period')
        context['lang'] = lang

        # attach translated titles for list
        page_obj = context.get("page_obj")
        if page_obj and lang != "ja":
            ids = [a.id for a in page_obj.object_list]
            tr_map = {
                t.article_id: t
                for t in ArticleTranslation.objects.filter(article_id__in=ids, language=lang)
            }
            for a in page_obj.object_list:
                tr = tr_map.get(a.id)
                a.display_title = tr.title_translated if tr else a.title
        else:
            for a in context.get("page_obj", []).object_list:
                a.display_title = a.title

        return context



@csrf_exempt  # Needed for cross-origin form posts
def receive_article(request):
    if request.method != 'POST':
        return JsonResponse({"error": "Only POST allowed"}, status=405)

    try:
        data = json.loads(request.body)
        title   = data.get('title') or ''
        text    = data.get('text') or ''
        url     = data.get('url') or ''
        publish = data.get('publish') or ''
        tag     = data.get('tag') or ''
        token   = data.get('secret_token')

        if token != 'TmeGoqJUSLcHelEpMdOeGKjw9hmBlgHMCF':
            return HttpResponseBadRequest("Invalid token.")

        User = get_user_model()
        user = User.objects.get(id=1)

        slug = slugify(title, allow_unicode=True)
        try:
            publish_dt = parse_japanese_date(publish)
        except Exception:
            publish_dt = now()

        # Try to find by slug OR url
        article = Article.objects.filter(Q(slug=slug) | Q(url=url)).first()
        created = False

        if not article:
            with transaction.atomic():
                # Re-check in case of race
                article = Article.objects.filter(Q(slug=slug) | Q(url=url)).select_for_update().first()
                if not article:
                    article = Article.objects.create(
                        title=title,
                        slug=slug,
                        url=url,
                        text=text,
                        publish=publish_dt,
                        user=user,
                    )
                    created = True

        if tag:
            article.tags.add(tag)

        return JsonResponse(
            {"message": "Article ready", "article_id": article.id, "created": created},
            status=201 if created else 200
        )
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


@csrf_exempt
def receive_translation(request):
    """
    Body: {article_id: int, language: 'en'|'zh', html: '<div id="engDiffBox">...</div>', secret_token: '...'}
    Stores:
      - title_translated = original article.title (trim to 300)
      - text_translated  = html (as-is)
    """
    if request.method != 'POST':
        return JsonResponse({"error": "Only POST allowed"}, status=405)

    try:
        data = json.loads(request.body)
        token      = data.get('secret_token')
        article_id = int(data.get('article_id'))
        language   = (data.get('language') or '').lower()
        html       = data.get('html') or ''

        if token != 'TmeGoqJUSLcHelEpMdOeGKjw9hmBlgHMCF':
            return HttpResponseBadRequest("Invalid token.")
        if language not in ('en', 'zh'):
            return HttpResponseBadRequest("language must be 'en' or 'zh'.")
        if not html:
            return HttpResponseBadRequest("html required.")

        article = Article.objects.get(id=article_id)

        obj, created = ArticleTranslation.objects.update_or_create(
            article=article, language=language,
            defaults=dict(
                # title_translated = None,
                text_translated  = html,  # store raw outerHTML
            )
        )
        return JsonResponse(
            {"message": f"translation {language} saved", "article_id": article.id, "created": created},
            status=201 if created else 200
        )

    except Article.DoesNotExist:
        return JsonResponse({"status": "error", "message": "article not found"}, status=404)
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)
    

def parse_japanese_date(date_str):
    """Parse date strings like '2025年8月6日 15:47' or '2025/8/8付'."""
    date_str = (date_str or "").strip()

    # Pattern 1: '2025年8月6日 15:47'
    match_full = re.match(r"(\d{4})年(\d{1,2})月(\d{1,2})日\s+(\d{1,2}):(\d{1,2})", date_str)
    if match_full:
        year, month, day, hour, minute = map(int, match_full.groups())
        return datetime(year, month, day, hour, minute)

    # Pattern 2: '2025/8/8付'
    match_date = re.match(r"(\d{4})/(\d{1,2})/(\d{1,2})", date_str)
    if match_date:
        year, month, day = map(int, match_date.groups())
        return datetime(year, month, day)

    return now()

# added for weekly news export
EXCLUDED_TAGS = {"unknown", "unknow"}   # case-insensitive
EXCERPT_CHARS = 180                     # JA excerpt length (~90 English words)

# --- Helpers ---
def excerpt_ja(text: str, n_chars: int = EXCERPT_CHARS) -> str:
    """Simple JA excerpt by characters; strips HTML & collapses whitespace."""
    s = strip_tags(text or "")
    s = " ".join(s.split())  # collapse all whitespace/newlines/tabs
    return s[:n_chars] + (" …" if len(s) > n_chars else "")

def clean_title(title: str) -> str:
    """Collapse whitespace so titles never include newlines/tabs."""
    return " ".join((title or "").split())

def derive_nikkei_translation_urls(url: str):
    """
    Build Nikkei translation URLs from the original article URL.
    Returns (en_url, zh_url) or (None, None) if not recognized.
    """
    if not url:
        return (None, None)
    try:
        u = urlparse(url)
    except Exception:
        return (None, None)

    if not u.netloc.endswith("nikkei.com"):
        return (None, None)

    # 1) from query ?ng=ID (already a translation link or similar)
    qs = parse_qs(u.query or "")
    article_id = qs.get("ng", [None])[0]

    # 2) from path /article/<ID>/
    if not article_id:
        m = re.search(r"/article/([A-Z0-9]+)/?", u.path or "")
        if m:
            article_id = m.group(1)

    # 3) last fallback: look for DGXZ*-style ID anywhere
    if not article_id:
        m = re.search(r"(DGXZ[A-Z0-9]+)", url)
        if m:
            article_id = m.group(1)

    if not article_id:
        return (None, None)

    base = f"https://www.nikkei.com/news/article-translation/?ng={article_id}"
    en = base
    zh = base + "&mta=c"
    return (en, zh)

def add_hyperlink(paragraph, url: str, text: str | None = None):
    """Insert a clickable hyperlink into a python-docx paragraph."""
    if not url:
        return
    from docx.oxml.shared import OxmlElement, qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    if text is None:
        text = url

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # underline + blue
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0000FF"); rPr.append(color)

    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    run.append(t)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)

def format_jp_range(start, end) -> str:
    """Format date range like '2025年7月21日～27日' (handles cross-month/year)."""
    if start.year == end.year:
        if start.month == end.month:
            return f"{start.year}年{start.month}月{start.day}日～{end.day}日"
        return f"{start.year}年{start.month}月{start.day}日～{end.month}月{end.day}日"
    return f"{start.year}年{start.month}月{start.day}日～{end.year}年{end.month}月{end.day}日"


# def weekly_news(request):
    """
    8-day JA weekly view:
      (1) list all tags for article.Article (sorted, excluding 'unknown'/'unknow'),
      (2) for each tag, fetch past-8-days JA news newest-first, or '今週主要なニュースなし',
      (3) HTML view and matching DOCX export (with clickable original / EN / ZH links).
    """
    # ---- Window: past 8 days including today ----
    today = timezone.localdate()
    start_date = today - timedelta(days=7)

    # ---- All tags for this content type (sorted by name), excluding unknown/unknow ----
    ct = ContentType.objects.get_for_model(Article)
    all_tags = (
        Tag.objects.filter(taggit_taggeditem_items__content_type=ct)
        .distinct()
    )

    def keep_tag(t: Tag) -> bool:
        return (t.name or "").strip().casefold() not in EXCLUDED_TAGS

    tags_sorted = sorted(
        (t for t in all_tags if keep_tag(t)),
        key=lambda t: (t.name or "").strip().casefold()
    )

    # ---- Build grouped data: one section per tag ----
    grouped = []  # list of dicts: {"tag_name": str, "articles": [Article]}
    for tag in tags_sorted:
        qs = (
            Article.objects
            .filter(
                language="ja",
                publish__range=(start_date, today),
                tags__in=[tag],
            )
            .order_by("-publish", "-created")
            .distinct()
        )
        arts = list(qs)

        # annotate each article
        for a in arts:
            a.display_title = clean_title(a.title)
            a.snippet = excerpt_ja(a.text)
            a.url_en, a.url_zh = derive_nikkei_translation_urls(a.url)

        grouped.append({"tag_name": tag.name, "articles": arts})

    # ---- DOCX export ----
    if request.GET.get("export") == "docx":
        from docx import Document
        from docx.shared import Pt  # <- for tight paragraph spacing

        doc = Document()

        # 3-line bold header with correct dates
        date_line = format_jp_range(start_date, today)
        for line in ("CP提携先企業動向まとめ", date_line, "日本正大光明 投資部"):
            p = doc.add_paragraph()
            p.add_run(line).bold = True

        for section in grouped:
            tag_name = section["tag_name"]
            arts = section["articles"]

            doc.add_heading(tag_name, level=2)

            if not arts:
                doc.add_paragraph("今週主要なニュースなし")
                continue

            for a in arts:
                # date (bold) + NEW line + title (bold)
                p = doc.add_paragraph()
                r1 = p.add_run(a.publish.strftime("%Y-%m-%d"))
                r1.add_break()
                p.add_run(a.display_title).bold = True
                p.paragraph_format.space_after = Pt(3)   #
                # excerpt with no extra spacing after
                if a.snippet:
                    p_snip = doc.add_paragraph(a.snippet)
                    p_snip.paragraph_format.space_after = Pt(0)   # ↓ tighten gap to next line

                # ONE LINE: 原文 (English translation / 中訳)
                if a.url or a.url_en or a.url_zh:
                    p_links = doc.add_paragraph()
                    fmt = p_links.paragraph_format
                    fmt.space_before = Pt(0)  # no extra gap above
                    fmt.space_after = Pt(0)   # no extra gap below

                    first = True
                    if a.url:
                        add_hyperlink(p_links, a.url, "原文")
                        first = False

                    # only add parentheses if we have at least one translation link
                    if a.url_en or a.url_zh:
                        p_links.add_run(" (")
                        wrote_any = False
                        if a.url_en:
                            add_hyperlink(p_links, a.url_en, "English translation")
                            wrote_any = True
                        if a.url_zh:
                            if wrote_any:
                                p_links.add_run(" / ")
                            add_hyperlink(p_links, a.url_zh, "中訳")
                        p_links.add_run(")")
                    p_links.paragraph_format.space_after = Pt(10)  # small gap after links

        buf = BytesIO()
        doc.save(buf); buf.seek(0)
        filename = f"weekly_news_ja_{today.strftime('%Y%m%d')}.docx"
        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp

    # ---- HTML render ----
    return render(request, "article/weekly_news.html", {
        "grouped": grouped,          # list of {"tag_name": str, "articles": [Article]}
        "start_date": start_date,
        "end_date": today,
    })
    
def weekly_news(request):
    """
    8-day JA weekly view (now with optional date window):
      (1) list all tags for article.Article (sorted, excluding 'unknown'/'unknow'),
      (2) for each tag, fetch news within [start_date, end_date] newest-first, or '今週主要なニュースなし',
      (3) HTML view and matching DOCX export (with clickable original / EN / ZH links).

    Optional query params:
      - ?start=YYYY-MM-DD
      - ?end=YYYY-MM-DD
      If neither provided, defaults to 'past 8 days including today'.
      If only one provided, expands to an 8-day window (inclusive).
    """
    # ---- Resolve window (defaults to past 8 days including today) ----
    today = timezone.localdate()

    def _parse_iso(dstr: str) -> date | None:
        try:
            return date.fromisoformat(dstr)
        except Exception:
            return None

    q_start = request.GET.get("start") or request.GET.get("start_date")
    q_end = request.GET.get("end") or request.GET.get("end_date")

    start_date = _parse_iso(q_start) if q_start else None
    end_date = _parse_iso(q_end) if q_end else None

    if start_date is None and end_date is None:
        # Default: past 8 days including today
        end_date = today
        start_date = end_date - timedelta(days=7)
    elif start_date is not None and end_date is None:
        # Expand to 8-day window from start
        end_date = start_date + timedelta(days=7)
    elif start_date is None and end_date is not None:
        # Expand to 8-day window ending at end
        start_date = end_date - timedelta(days=7)
    # else: both provided → use as-is

    # If user accidentally swapped them, fix.
    if start_date > end_date:
        start_date, end_date = end_date, start_date

    # ---- All tags for this content type (sorted by name), excluding unknown/unknow ----
    ct = ContentType.objects.get_for_model(Article)
    all_tags = (
        Tag.objects.filter(taggit_taggeditem_items__content_type=ct)
        .distinct()
    )

    def keep_tag(t: Tag) -> bool:
        return (t.name or "").strip().casefold() not in EXCLUDED_TAGS

    tags_sorted = sorted(
        (t for t in all_tags if keep_tag(t)),
        key=lambda t: (t.name or "").strip().casefold()
    )

    # ---- Build grouped data: one section per tag ----
    grouped = []  # list of dicts: {"tag_name": str, "articles": [Article]}
    for tag in tags_sorted:
        qs = (
            Article.objects
            .filter(
                language="ja",
                publish__range=(start_date, end_date),
                tags__in=[tag],
            )
            .order_by("-publish", "-created")
            .distinct()
        )
        arts = list(qs)

        # annotate each article
        for a in arts:
            a.display_title = clean_title(a.title)
            a.snippet = excerpt_ja(a.text)
            a.url_en, a.url_zh = derive_nikkei_translation_urls(a.url)

        grouped.append({"tag_name": tag.name, "articles": arts})

    # ---- DOCX export ----
    if request.GET.get("export") == "docx":
        from docx import Document
        from docx.shared import Pt  # <- for tight paragraph spacing
        from io import BytesIO
        from django.http import HttpResponse

        doc = Document()

        # 3-line bold header with selected dates
        date_line = format_jp_range(start_date, end_date)
        for line in ("CP提携先企業動向まとめ", date_line, "日本正大光明 投資部"):
            p = doc.add_paragraph()
            p.add_run(line).bold = True

        for section in grouped:
            tag_name = section["tag_name"]
            arts = section["articles"]

            doc.add_heading(tag_name, level=2)

            if not arts:
                doc.add_paragraph("今週主要なニュースなし")
                continue

            for a in arts:
                # date (bold) + NEW line + title (bold)
                p = doc.add_paragraph()
                r1 = p.add_run(a.publish.strftime("%Y-%m-%d"))
                r1.add_break()
                p.add_run(a.display_title).bold = True
                p.paragraph_format.space_after = Pt(3)
                # excerpt with no extra spacing after
                if a.snippet:
                    p_snip = doc.add_paragraph(a.snippet)
                    p_snip.paragraph_format.space_after = Pt(0)

                # ONE LINE: 原文 (English translation / 中訳)
                if a.url or a.url_en or a.url_zh:
                    p_links = doc.add_paragraph()
                    fmt = p_links.paragraph_format
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)

                    first = True
                    if a.url:
                        add_hyperlink(p_links, a.url, "原文")
                        first = False

                    if a.url_en or a.url_zh:
                        p_links.add_run(" (")
                        wrote_any = False
                        if a.url_en:
                            add_hyperlink(p_links, a.url_en, "English translation")
                            wrote_any = True
                        if a.url_zh:
                            if wrote_any:
                                p_links.add_run(" / ")
                            add_hyperlink(p_links, a.url_zh, "中訳")
                        p_links.add_run(")")
                    p_links.paragraph_format.space_after = Pt(10)

        buf = BytesIO()
        doc.save(buf); buf.seek(0)
        filename = f"weekly_news_ja_{end_date.strftime('%Y%m%d')}.docx"
        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename=\"{filename}\"'
        return resp
